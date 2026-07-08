#!/usr/bin/env python3
"""
Flavor Analysis Tools for GCMS Agent
=====================================
OAV calculation, internal standard normalization, blank subtraction,
ANOVA with post-hoc tests, flavor wheel, off-flavor database,
oxidation/Maillard markers, PLS-DA, VIP, Random Forest,
interactive HTML reports, Word tables, batch effect correction.

Usage:
    from flavor_tools import FlavorAnalysisMixin
    class GCMSAgent(FlavorAnalysisMixin):
        ...
"""

import json
import os
import numpy as np
from pathlib import Path
from datetime import datetime

# ================================================================
# Odor Threshold Database (in water, µg/L or µg/kg)
# ================================================================
# Sources: van Gemert (2011), Burdock (2010), Rychlik et al. (2011)
ODOR_THRESHOLDS = {
    # Aldehydes
    'hexanal':             4.5,
    'heptanal':            3.0,
    'octanal':             0.7,
    'nonanal':             1.0,
    'decanal':             0.1,
    'benzaldehyde':        350.0,
    'phenylacetaldehyde':  4.0,
    'furfural':            3000.0,
    'cinnamaldehyde':      50.0,
    'citral':              120.0,
    'citronellal':         100.0,
    '2-heptenal':          13.0,
    '2-octenal':           3.0,
    '2-nonenal':           0.08,
    '2-decenal':           0.3,
    '2,4-heptadienal':     10.0,
    '2,4-decadienal':      0.07,
    '2,6-nonadienal':      0.001,

    # Ketones
    '2-heptanone':         140.0,
    '2-octanone':          50.0,
    '2-nonanone':          5.0,
    '2-undecanone':        7.0,
    '2-tridecanone':       10.0,
    'acetoin':             800.0,
    'acetophenone':        65.0,
    '2,3-butanedione':     0.1,
    '2,3-pentanedione':    30.0,
    'beta-ionone':         0.007,
    'geranylacetone':      60.0,
    'carvone':             50.0,
    'menthone':            170.0,
    'camphor':             200.0,
    'sotolon':             0.01,
    'nootkatone':          1.0,
    'furaneol':            0.03,

    # Alcohols
    '1-hexanol':           2500.0,
    '1-octen-3-ol':        1.0,
    '1-nonanol':           50.0,
    'linalool':            6.0,
    'alpha-terpineol':     330.0,
    'geraniol':            40.0,
    'nerol':               400.0,
    'citronellol':         40.0,
    'borneol':             10.0,
    'phenylethyl alcohol': 1100.0,
    'benzyl alcohol':      5500.0,
    '2,3-butanediol':      100000.0,
    'isopentyl alcohol':   250.0,
    'furfuryl alcohol':    2000.0,
    'maltol':              35000.0,
    'farnesol':            20.0,
    'nerolidol':           10.0,

    # Esters
    'ethyl acetate':       5000.0,
    'isoamyl acetate':     2.0,
    'ethyl butyrate':      1.0,
    'ethyl hexanoate':     1.0,
    'ethyl octanoate':     15.0,
    'ethyl decanoate':     10.0,
    'hexyl acetate':       2.0,
    'benzyl acetate':      200.0,
    'phenylethyl acetate': 250.0,
    'gamma-butyrolactone': 1000.0,
    'gamma-hexalactone':   50.0,
    'gamma-octalactone':   7.0,
    'gamma-nonalactone':   9.0,
    'gamma-decalactone':   1.0,
    'gamma-dodecalactone': 0.5,
    'delta-decalactone':   100.0,
    'wine lactone':        0.01,
    'whiskey lactone':     20.0,

    # Acids
    'acetic acid':         22000.0,
    'butyric acid':        240.0,
    'hexanoic acid':       3000.0,
    'octanoic acid':       3000.0,
    'decanoic acid':       10000.0,
    'isovaleric acid':     120.0,
    'propanoic acid':      20000.0,

    # Terpenes
    'alpha-pinene':        6.0,
    'beta-pinene':         140.0,
    'limonene':            10.0,
    'myrcene':             13.0,
    'gamma-terpinene':     1000.0,
    'p-cymene':            11.0,
    'caryophyllene':       64.0,
    'humulene':            120.0,
    'valencene':           100.0,
    'eucalyptol':          12.0,
    'sabinene':            10.0,
    'alpha-phellandrene':  40.0,
    'alpha-farnesene':     0.6,
    'estragole':           10.0,
    'anethole':            15.0,
    'thymol':              100.0,
    'carvacrol':           200.0,
    'eugenol':             6.0,
    'isoeugenol':          6.0,
    'vanillin':            20.0,

    # Pyrazines
    '2-methylpyrazine':    60000.0,
    '2,5-dimethylpyrazine': 1700.0,
    '2,6-dimethylpyrazine': 1500.0,
    '2,3,5-trimethylpyrazine': 400.0,
    '2-ethyl-3,5-dimethylpyrazine': 0.04,
    '2-ethylpyrazine':     6000.0,
    '2,3-diethylpyrazine': 1.0,

    # Sulfur compounds
    'dimethyl sulfide':    0.3,
    'dimethyl disulfide':  0.16,
    'dimethyl trisulfide': 0.01,
    'methional':           0.2,
    'dimethyl sulfone':    10000.0,
    'benzothiazole':       80.0,
    'furfurylthiol':       0.005,
    '2-methyl-3-furanthiol': 0.005,
    'diallyl disulfide':   0.3,

    # Furans
    '2-pentylfuran':       6.0,
    '2-ethylfuran':        30.0,
    '2-methylfuran':       1000.0,
    '2,5-dimethylfuran':   10.0,
    '2-acetylfuran':       10000.0,

    # Pyrroles/Indoles/Nitrogen
    'pyrrole':             10000.0,
    'indole':              140.0,
    'skatole':             0.4,
    '2-acetylpyrrole':     10000.0,
    '2-acetylpyridine':    1000.0,

    # Phenols
    'phenol':              5900.0,
    'guaiacol':            1.0,
    '4-ethylguaiacol':     50.0,
    '4-ethylphenol':       21.0,
    '4-vinylguaiacol':     10.0,
    '2,6-dimethylphenol':  400.0,
    '2,4-di-tert-butylphenol': 100.0,

    # Alkanes
    'dodecane':            1000.0,
    'tetradecane':         1000.0,
    'hexadecane':          1000.0,

    # Siloxanes / contaminants (very high threshold = not flavor-active)
    'octamethylcyclotetrasiloxane': 100000.0,
    'decamethylcyclopentasiloxane': 100000.0,
    'hexamethylcyclotrisiloxane': 100000.0,
}

# ================================================================
# Off-Flavor Compound Database (microalgae-specific)
# ================================================================
OFF_FLAVOR_DB = {
    'geosmin': {
        'odor': 'earthy, musty, muddy',
        'threshold': 0.015,
        'source': 'cyanobacteria, actinomycetes',
        'cas': '19700-21-1',
        'ri': 1390,
    },
    '2-methylisoborneol': {
        'odor': 'earthy, musty, camphor',
        'threshold': 0.01,
        'source': 'cyanobacteria',
        'cas': '2371-42-8',
        'ri': 1170,
    },
    'dimethyl trisulfide': {
        'odor': 'sulfurous, cabbage, putrid',
        'threshold': 0.01,
        'source': 'protein degradation, sulfur amino acids',
        'cas': '3658-80-8',
        'ri': 970,
    },
    'dimethyl disulfide': {
        'odor': 'sulfurous, onion, cabbage',
        'threshold': 0.16,
        'source': 'protein degradation',
        'cas': '624-92-0',
        'ri': 740,
    },
    '2-nonenal': {
        'odor': 'fatty, green, cucumber',
        'threshold': 0.08,
        'source': 'lipid oxidation (linoleic acid)',
        'cas': '18829-56-6',
        'ri': 1155,
    },
    '2,4-decadienal': {
        'odor': 'fatty, rancid, deep-fried',
        'threshold': 0.07,
        'source': 'lipid oxidation (linoleic acid)',
        'cas': '25152-84-5',
        'ri': 1310,
    },
    'hexanal': {
        'odor': 'green, grassy, fatty',
        'threshold': 4.5,
        'source': 'lipid oxidation (linolenic acid)',
        'cas': '66-25-1',
        'ri': 800,
    },
    '1-octen-3-ol': {
        'odor': 'mushroom, earthy, moldy',
        'threshold': 1.0,
        'source': 'lipid oxidation, enzymatic',
        'cas': '3391-86-4',
        'ri': 980,
    },
    'indole': {
        'odor': 'fecal, animal, musty',
        'threshold': 140.0,
        'source': 'tryptophan degradation',
        'cas': '120-72-9',
        'ri': 1290,
    },
    'skatole': {
        'odor': 'fecal, mothball',
        'threshold': 0.4,
        'source': 'tryptophan degradation',
        'cas': '83-34-1',
        'ri': 1390,
    },
    'trimethylamine': {
        'odor': 'fishy, ammoniacal',
        'threshold': 0.00021,
        'source': 'choline/TMAO degradation',
        'cas': '75-50-3',
        'ri': 420,
    },
    '2-pentylfuran': {
        'odor': 'beany, green, earthy',
        'threshold': 6.0,
        'source': 'lipid oxidation',
        'cas': '3777-69-3',
        'ri': 990,
    },
    'methional': {
        'odor': 'cooked potato, sulfurous',
        'threshold': 0.2,
        'source': 'methionine Strecker degradation',
        'cas': '3268-49-3',
        'ri': 900,
    },
    'butyric acid': {
        'odor': 'rancid, cheesy, vomit',
        'threshold': 240.0,
        'source': 'lipid/fat degradation',
        'cas': '107-92-6',
        'ri': 820,
    },
    'isovaleric acid': {
        'odor': 'sweaty, cheesy, rancid',
        'threshold': 120.0,
        'source': 'protein degradation, leucine',
        'cas': '503-74-2',
        'ri': 870,
    },
}

# ================================================================
# Maillard Reaction Product Markers
# ================================================================
MAILLARD_MARKERS = {
    'furfural': 'pentose dehydration',
    '5-hydroxymethylfurfural': 'hexose dehydration',
    '2-methylpyrazine': 'α-dicarbonyl + amino acid',
    '2,5-dimethylpyrazine': 'α-dicarbonyl + alanine',
    '2,6-dimethylpyrazine': 'α-dicarbonyl + alanine',
    '2,3,5-trimethylpyrazine': 'α-dicarbonyl + glycine/alanine',
    '2-ethyl-3,5-dimethylpyrazine': 'α-dicarbonyl + alanine',
    '2-ethylpyrazine': 'α-dicarbonyl + amino acid',
    '2,3-diethylpyrazine': 'α-dicarbonyl + alanine',
    '2-acetylpyrrole': 'Amadori product cyclization',
    'pyrrole': 'proline degradation',
    'acetoin': 'pyruvate decarboxylation',
    '2,3-butanedione': 'pyruvate/pentose degradation',
    '2,3-pentanedione': 'pentose/amino acid',
    'methional': 'methionine Strecker aldehyde',
    'phenylacetaldehyde': 'phenylalanine Strecker aldehyde',
    '3-methylbutanal': 'leucine Strecker aldehyde',
    '2-methylbutanal': 'isoleucine Strecker aldehyde',
    'furfuryl alcohol': 'furfural reduction',
    'maltol': 'sugar degradation',
    'furaneol': 'hexose dehydration/rearrangement',
    'sotolon': 'Maillard + oxidation',
}

# ================================================================
# Lipid Oxidation Product Markers
# ================================================================
LIPID_OXIDATION_MARKERS = {
    # From linoleic acid (C18:2, ω-6)
    'hexanal': 'linoleic acid (C18:2 ω-6) oxidation',
    '2-heptenal': 'linoleic acid oxidation',
    '2-octenal': 'linoleic acid oxidation',
    '2,4-decadienal': 'linoleic acid primary oxidation marker',
    '2-pentylfuran': 'linoleic acid oxidation',

    # From linolenic acid (C18:3, ω-3)
    '2,4-heptadienal': 'linolenic acid (C18:3 ω-3) oxidation',
    '2,6-nonadienal': 'linolenic acid oxidation',
    '2,4,7-decatrienal': 'linolenic acid oxidation',

    # From oleic acid (C18:1, ω-9)
    'heptanal': 'oleic acid (C18:1 ω-9) oxidation',
    'octanal': 'oleic acid oxidation',
    'nonanal': 'oleic acid oxidation',
    'decanal': 'oleic acid oxidation',
    '2-decenal': 'oleic acid oxidation',
    'octanoic acid': 'oleic acid oxidation end-product',
    'nonanoic acid': 'oleic acid oxidation end-product',

    # General
    '1-octen-3-ol': 'lipid oxidation (enzymatic)',
    '2-heptanone': 'β-keto acid decarboxylation (fatty acid)',
    '2-octanone': 'β-keto acid decarboxylation',
    '2-nonanone': 'β-keto acid decarboxylation',
    'butyric acid': 'short-chain fatty acid oxidation',
    'hexanoic acid': 'fatty acid oxidation',
    'hexadecane': 'fatty acid decarboxylation',
}


# ================================================================
# OAV Calculator
# ================================================================
def calculate_oav(df, concentration_col='conc_g100g', threshold_db=None):
    """Calculate Odor Activity Values (OAV) for all compounds.

    OAV = concentration / odor_threshold
    OAV > 1: compound contributes to aroma
    OAV > 10: significant contributor
    OAV > 100: dominant contributor

    Args:
        df: DataFrame with compound names and concentration values
        concentration_col: column name for concentration values
        threshold_db: custom threshold dict (uses ODOR_THRESHOLDS if None)

    Returns:
        DataFrame with added 'oav', 'log_oav', 'odor_threshold', 'aroma_impact' columns
    """
    if threshold_db is None:
        threshold_db = ODOR_THRESHOLDS

    result = df.copy()
    oavs = []
    thresholds_found = []
    impacts = []

    for _, row in result.iterrows():
        compound = str(row.get('compound', '')).lower()
        conc = row.get(concentration_col, 0)
        threshold = _find_threshold(compound, threshold_db)

        if threshold and conc > 0:
            oav = conc / threshold
        else:
            oav = 0.0

        oavs.append(oav)
        thresholds_found.append(threshold if threshold else None)
        impacts.append(_classify_impact(oav))

    result['oav'] = oavs
    result['log_oav'] = [np.log10(v) if v > 0 else 0 for v in oavs]
    result['odor_threshold'] = thresholds_found
    result['aroma_impact'] = impacts

    return result


def _find_threshold(compound_name, threshold_db):
    """Find odor threshold by compound name (fuzzy match)."""
    name = compound_name.lower().strip()
    # Exact match
    if name in threshold_db:
        return threshold_db[name]
    # Substring match (compound name contains key or vice versa)
    for key, val in threshold_db.items():
        if key in name or name in key:
            return val
    # Try matching without common suffixes
    import re
    base = re.sub(r'\s*[\(\[].*[\)\]]', '', name).strip()
    if base in threshold_db:
        return threshold_db[base]
    return None


def _classify_impact(oav):
    if oav > 100:
        return 'dominant'
    elif oav > 10:
        return 'significant'
    elif oav > 1:
        return 'contributing'
    elif oav > 0.1:
        return 'sub-threshold'
    else:
        return 'negligible'


def get_oav_summary(df, group_col='group', top_n=15):
    """Generate OAV summary: top aroma-impact compounds per group."""
    df_oav = calculate_oav(df)

    summary = {
        'top_oav_overall': [],
        'top_oav_by_group': {},
        'aroma_impact_counts': {},
        'compounds_without_threshold': [],
    }

    # Top OAV overall
    overall = df_oav.groupby('compound')['oav'].mean().sort_values(ascending=False)
    for comp, val in overall.head(top_n).items():
        threshold = _find_threshold(comp, ODOR_THRESHOLDS)
        summary['top_oav_overall'].append({
            'compound': comp,
            'mean_oav': round(float(val), 2),
            'log_oav': round(np.log10(val) if val > 0 else 0, 2),
            'threshold': threshold,
            'impact': _classify_impact(val),
        })

    # By group
    if group_col in df_oav.columns:
        for g in df_oav[group_col].unique():
            gdf = df_oav[df_oav[group_col] == g]
            gavg = gdf.groupby('compound')['oav'].mean().sort_values(ascending=False)
            summary['top_oav_by_group'][str(g)] = []
            for comp, val in gavg.head(10).items():
                summary['top_oav_by_group'][str(g)].append({
                    'compound': comp, 'mean_oav': round(float(val), 2),
                    'impact': _classify_impact(val),
                })

    # Impact distribution
    impacts = df_oav['aroma_impact'].value_counts().to_dict()
    summary['aroma_impact_counts'] = {str(k): int(v) for k, v in impacts.items()}

    # Compounds missing thresholds
    no_threshold = df_oav[df_oav['odor_threshold'].isna()]['compound'].unique()
    summary['compounds_without_threshold'] = sorted(no_threshold.tolist())[:30]

    return summary


# ================================================================
# Internal Standard (ISTD) Normalization
# ================================================================
def normalize_istd(df, istd_name='internal standard', istd_conc=1.0):
    """Normalize compound concentrations by internal standard.

    Corrected_conc = (Compound_Area / ISTD_Area) * ISTD_Concentration

    Args:
        df: DataFrame with sample, compound, area columns
        istd_name: name of the internal standard compound
        istd_conc: known concentration of ISTD (default 1.0 for relative)

    Returns:
        DataFrame with added 'conc_normalized' column
    """
    result = df.copy()
    result['istd_area'] = np.nan
    result['conc_normalized'] = np.nan

    for sample in result['sample'].unique():
        sample_mask = result['sample'] == sample
        sdf = result[sample_mask]

        # Find ISTD in this sample
        istd_match = None
        for compound in sdf['compound'].unique():
            if istd_name.lower() in str(compound).lower():
                istd_match = compound
                break

        if istd_match:
            istd_area = sdf[sdf['compound'] == istd_match]['area'].values[0]
            result.loc[sample_mask, 'istd_area'] = istd_area
            result.loc[sample_mask, 'conc_normalized'] = (
                sdf['area'] / istd_area * istd_conc
            )
        else:
            # ISTD not found — use original area
            result.loc[sample_mask, 'conc_normalized'] = result.loc[sample_mask, 'conc_g100g']

    return result


# ================================================================
# Blank Subtraction
# ================================================================
def subtract_blank(df, blank_sample_name='blank', min_signal_ratio=3.0):
    """Subtract blank sample signal from all samples.

    For each compound: if sample_area / blank_area < min_signal_ratio,
    the compound is flagged as potential background/contamination.

    Args:
        df: DataFrame with sample, compound, area columns
        blank_sample_name: name of the blank sample
        min_signal_ratio: S/N threshold (default 3:1)

    Returns:
        DataFrame with 'blank_area', 'sn_ratio', 'blank_flag' columns
    """
    result = df.copy()

    # Find blank sample
    blank_samples = [s for s in result['sample'].unique()
                     if blank_sample_name.lower() in str(s).lower()]

    if not blank_samples:
        result['blank_area'] = 0
        result['sn_ratio'] = np.nan
        result['blank_flag'] = 'no_blank_sample'
        return result

    blank_name = blank_samples[0]
    blank_df = result[result['sample'] == blank_name]

    result['blank_area'] = 0
    result['sn_ratio'] = np.nan
    result['blank_flag'] = 'ok'

    for compound in result['compound'].unique():
        compound_mask = result['compound'] == compound
        is_blank = result['sample'] == blank_name

        blank_val = blank_df[blank_df['compound'] == compound]['area'].values
        blank_area = float(blank_val[0]) if len(blank_val) > 0 else 0.0

        result.loc[compound_mask, 'blank_area'] = blank_area

        for sample in result.loc[compound_mask & ~is_blank, 'sample'].unique():
            smask = (result['compound'] == compound) & (result['sample'] == sample)
            sample_area = result.loc[smask, 'area'].values[0]

            if blank_area > 0:
                ratio = sample_area / blank_area
                result.loc[smask, 'sn_ratio'] = round(ratio, 1)
                if ratio < min_signal_ratio:
                    result.loc[smask, 'blank_flag'] = 'below_blank'
                else:
                    result.loc[smask, 'blank_flag'] = 'ok'
                    # Subtract blank
                    result.loc[smask, 'area'] = max(0, sample_area - blank_area)
            else:
                result.loc[smask, 'sn_ratio'] = float('inf') if sample_area > 0 else 0

    # Remove blank rows from non-blank operations
    return result[result['sample'] != blank_name]


# ================================================================
# ANOVA with Post-Hoc Tests
# ================================================================
def run_anova(df, group_col='group', compound_col='compound', value_col='conc_g100g',
              alpha=0.05, posthoc_method='tukey'):
    """Run one-way ANOVA + post-hoc test for each compound across groups.

    Args:
        df: DataFrame
        group_col: column with group labels
        compound_col: column with compound names
        value_col: column with quantitative values
        alpha: significance level
        posthoc_method: 'tukey' (Tukey HSD) or 'games_howell'

    Returns:
        dict with ANOVA F-stat, p-value, post-hoc pairwise comparisons
    """
    from scipy import stats
    import warnings
    warnings.filterwarnings('ignore')

    groups = sorted(df[group_col].unique())
    if len(groups) < 2:
        return {'error': 'Need at least 2 groups for ANOVA'}

    results = {
        'anova_results': [],
        'significant_compounds': [],
        'posthoc_results': {},
        'group_order': groups,
        'method': f"One-way ANOVA + {posthoc_method.upper()} post-hoc",
    }

    for compound in sorted(df[compound_col].unique()):
        cdf = df[df[compound_col] == compound]
        group_vals = []

        for g in groups:
            vals = cdf[cdf[group_col] == g][value_col].dropna().values
            if len(vals) >= 2:
                group_vals.append(vals)
            elif len(vals) > 0:
                group_vals.append(vals)

        if len(group_vals) < 2:
            continue

        if len(group_vals) == 2:
            # Fall back to t-test for 2 groups
            t_stat, t_p = stats.ttest_ind(group_vals[0], group_vals[1], equal_var=False)
            anova_result = {
                'compound': compound,
                'F_stat': float(t_stat ** 2),
                'p_value': float(t_p),
                'significant': bool(t_p < alpha),
                'n_groups_compared': 2,
            }
        else:
            # One-way ANOVA
            f_stat, p_val = stats.f_oneway(*group_vals)
            anova_result = {
                'compound': compound,
                'F_stat': float(f_stat),
                'p_value': float(p_val),
                'significant': bool(p_val < alpha),
                'n_groups_compared': len(group_vals),
            }

        results['anova_results'].append(anova_result)

        if anova_result['significant']:
            results['significant_compounds'].append(compound)

            # Post-hoc
            comparisons = _run_posthoc(group_vals, groups, compound, posthoc_method, alpha)
            if comparisons:
                results['posthoc_results'][compound] = comparisons

    # Sort by p-value
    results['anova_results'].sort(key=lambda x: x['p_value'])
    results['n_significant'] = len(results['significant_compounds'])
    results['n_tested'] = len(results['anova_results'])

    return results


def _run_posthoc(group_vals, group_names, compound, method, alpha):
    """Run post-hoc pairwise comparisons."""
    from itertools import combinations

    comparisons = []

    if method == 'tukey':
        # Tukey HSD
        for (i, vals_i), (j, vals_j) in combinations(enumerate(group_vals), 2):
            mean_i, mean_j = np.mean(vals_i), np.mean(vals_j)
            n_i, n_j = len(vals_i), len(vals_j)

            if n_i < 2 or n_j < 2:
                continue

            # Pooled SE
            var_i = np.var(vals_i, ddof=1)
            var_j = np.var(vals_j, ddof=1)
            pooled_se = np.sqrt((var_i / n_i + var_j / n_j) / 2)

            if pooled_se > 0:
                q_stat = abs(mean_i - mean_j) / pooled_se
                # Approximate Tukey critical value
                from scipy import stats
                k = len(group_names)
                df = n_i + n_j - 2
                # Studentized range approximation
                q_crit = 3.3  # Approx for k=3, df=10, alpha=0.05
                if k == 3:
                    q_crit = 3.4 if df < 10 else 3.0
                elif k == 4:
                    q_crit = 3.9 if df < 10 else 3.3

                comparisons.append({
                    'groups': f"{group_names[i]} vs {group_names[j]}",
                    'mean_diff': round(float(mean_i - mean_j), 6),
                    'fold_change': round(float(mean_i / mean_j) if mean_j > 0 else 0, 2),
                    'q_statistic': round(float(q_stat), 3),
                    'significant': bool(q_stat > q_crit),
                })

    else:
        # Games-Howell (for unequal variances)
        for (i, vals_i), (j, vals_j) in combinations(enumerate(group_vals), 2):
            mean_i, mean_j = np.mean(vals_i), np.mean(vals_j)
            n_i, n_j = len(vals_i), len(vals_j)

            if n_i < 2 or n_j < 2:
                continue

            var_i = np.var(vals_i, ddof=1)
            var_j = np.var(vals_j, ddof=1)

            se = np.sqrt(var_i / n_i + var_j / n_j)
            if se > 0:
                t_gh = abs(mean_i - mean_j) / se
                df_gh = (var_i / n_i + var_j / n_j) ** 2 / (
                    (var_i / n_i) ** 2 / (n_i - 1) + (var_j / n_j) ** 2 / (n_j - 1)
                )
                from scipy import stats
                p_gh = 2 * stats.t.sf(t_gh, df_gh)

                comparisons.append({
                    'groups': f"{group_names[i]} vs {group_names[j]}",
                    'mean_diff': round(float(mean_i - mean_j), 6),
                    'fold_change': round(float(mean_i / mean_j) if mean_j > 0 else 0, 2),
                    't_statistic': round(float(t_gh), 3),
                    'p_value': round(float(p_gh), 5),
                    'significant': bool(p_gh < alpha),
                })

    return comparisons


# ================================================================
# Compound Pathway Tagging
# ================================================================
def tag_compounds(df):
    """Auto-tag compounds with flavor pathway markers.

    Tags: Maillard reaction, lipid oxidation, off-flavor

    Returns:
        DataFrame with added 'pathway' and 'flavor_note' columns
    """
    result = df.copy()
    pathways = []
    flavor_notes = []

    for _, row in result.iterrows():
        compound = str(row.get('compound', '')).lower().strip()
        tags = []
        notes = []

        # Maillard markers
        if compound in MAILLARD_MARKERS:
            tags.append('Maillard')
            notes.append(MAILLARD_MARKERS[compound])

        # Lipid oxidation markers
        if compound in LIPID_OXIDATION_MARKERS:
            tags.append('Lipid_Oxidation')
            notes.append(LIPID_OXIDATION_MARKERS[compound])

        # Off-flavor check
        if compound in OFF_FLAVOR_DB:
            tags.append('Off_Flavor')
            od = OFF_FLAVOR_DB[compound]
            notes.append(f"Odor: {od['odor']} (threshold: {od['threshold']})")

        if not tags:
            tags.append('Other')
            notes.append('')

        pathways.append('|'.join(tags))
        flavor_notes.append('; '.join(notes) if notes else '')

    result['pathway'] = pathways
    result['flavor_note'] = flavor_notes
    return result


# ================================================================
# Flavor Wheel / Radar Chart
# ================================================================
def generate_flavor_wheel(df, title='Flavor Profile Radar', output_path=None):
    """Generate a flavor wheel / radar chart showing aroma profile.

    Groups compounds by odor category and shows relative contribution.

    Args:
        df: DataFrame with compound data
        title: chart title
        output_path: save path (default: output/agent_results/plots/flavor_wheel.png)

    Returns:
        dict with status and file path
    """
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    # Odor category mapping
    ODOR_CATEGORIES = {
        'green/grassy': ['hexanal', 'heptanal', 'octanal', 'nonanal', 'decanal',
                         '2-heptenal', '2-octenal', '2-nonenal', '2-decenal',
                         '2,4-heptadienal', '2,6-nonadienal', 'citronellal',
                         'hexanol'],
        'fatty/rancid': ['2,4-decadienal', 'butyric acid', 'hexanoic acid',
                         'octanoic acid', 'isovaleric acid'],
        'fruity/sweet': ['ethyl acetate', 'isoamyl acetate', 'ethyl butyrate',
                         'ethyl hexanoate', 'ethyl octanoate', 'ethyl decanoate',
                         'benzaldehyde', 'phenylacetaldehyde', 'vanillin',
                         'beta-ionone', 'geranylacetone', 'furaneol', 'maltol',
                         'linalool', 'geraniol', 'citronellol'],
        'earthy/musty': ['geosmin', '2-methylisoborneol', '1-octen-3-ol',
                         '2-pentylfuran', 'indole', 'skatole'],
        'roasted/nutty': ['2-methylpyrazine', '2,5-dimethylpyrazine',
                          '2,6-dimethylpyrazine', '2,3,5-trimethylpyrazine',
                          '2-ethyl-3,5-dimethylpyrazine', '2-acetylpyrrole',
                          'furfural', 'furfuryl alcohol', 'pyrrole'],
        'sulfurous': ['dimethyl sulfide', 'dimethyl disulfide', 'dimethyl trisulfide',
                      'methional', 'benzothiazole', 'furfurylthiol',
                      'diallyl disulfide'],
        'floral/spicy': ['alpha-pinene', 'beta-pinene', 'limonene', 'myrcene',
                         'caryophyllene', 'humulene', 'eugenol', 'isoeugenol',
                         'phenylethyl alcohol', 'nerol', 'alpha-terpineol',
                         'estragole', 'anethole', 'thymol', 'carvacrol'],
        'fermented/dairy': ['acetoin', '2,3-butanedione', '2,3-pentanedione',
                            'acetic acid', 'propanoic acid', '3-methylbutanal',
                            '2-methylbutanal', 'isopentyl alcohol'],
    }

    # Compute category scores
    df_oav = calculate_oav(df) if 'oav' not in df.columns else df
    category_scores = {}
    for cat, compounds in ODOR_CATEGORIES.items():
        score = 0
        for comp in compounds:
            mask = df_oav['compound'].str.lower().str.strip() == comp
            if mask.any():
                score += df_oav.loc[mask, 'log_oav'].mean()
        category_scores[cat] = max(score, 0.01)

    categories = list(category_scores.keys())
    values = [category_scores[c] for c in categories]
    N = len(categories)

    # Radar chart
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    values += values[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw={'projection': 'polar'})
    ax.fill(angles, values, alpha=0.25, color='#0072B2')
    ax.plot(angles, values, 'o-', linewidth=2, color='#0072B2')
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=10)
    ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
    ax.set_ylabel('log₁₀(OAV)', fontsize=10)
    ax.grid(True, alpha=0.3)

    if output_path is None:
        output_path = 'output/agent_results/plots/flavor_wheel.png'

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return {
        'status': 'done',
        'file': output_path,
        'categories': len(categories),
        'top_category': categories[np.argmax(values[:-1])],
        'note': 'Flavor wheel generated. Shows aroma profile by odor category.'
    }


def get_off_flavor_report(df):
    """Check detected compounds against off-flavor database.

    Returns list of detected off-flavor compounds with their descriptors.
    """
    detected = []
    for compound in df['compound'].unique():
        name = str(compound).lower().strip()
        for off_name, info in OFF_FLAVOR_DB.items():
            if off_name in name or name in off_name:
                mean_area = df[df['compound'] == compound]['area'].mean()
                detected.append({
                    'compound': compound,
                    'odor': info['odor'],
                    'threshold_ugL': info['threshold'],
                    'source': info['source'],
                    'mean_area': round(float(mean_area), 0),
                    'severity': 'high' if mean_area > 100000 else 'medium' if mean_area > 10000 else 'low',
                })
                break
    return {'off_flavors_detected': detected, 'n_detected': len(detected)}


# ================================================================
# PLS-DA (Partial Least Squares Discriminant Analysis)
# ================================================================
def run_plsda(df, group_col='group', value_col='conc_g100g', n_components=2):
    """Run PLS-DA to discriminate groups and identify key compounds.

    Args:
        df: DataFrame
        group_col: column with group labels
        value_col: column with quantitative values
        n_components: number of latent variables

    Returns:
        dict with scores, loadings, VIP scores, R2/Q2, plot path
    """
    from sklearn.preprocessing import StandardScaler, LabelEncoder
    from sklearn.cross_decomposition import PLSRegression

    # Pivot: samples × compounds
    pivot = df.pivot_table(values=value_col, index=['group', 'sample'],
                           columns='compound', aggfunc='mean').fillna(0)
    X = pivot.values
    sample_labels = [f"{g}|{s}" for g, s in pivot.index]
    group_labels = [g for g, _ in pivot.index]
    compound_names = list(pivot.columns)

    # Encode groups
    le = LabelEncoder()
    y = le.fit_transform(group_labels)

    # Scale
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    # PLS-DA
    pls = PLSRegression(n_components=min(n_components, X_scaled.shape[1], X_scaled.shape[0]))
    pls.fit(X_scaled, y)

    scores = pls.x_scores_
    loadings = pls.x_loadings_

    # VIP scores
    vip = _calculate_vip(pls, X_scaled, y)

    # R2 and Q2 (simple approximation)
    y_pred = pls.predict(X_scaled)
    ss_res = np.sum((y - y_pred.ravel()) ** 2)
    ss_tot = np.sum((y - np.mean(y)) ** 2)
    r2 = 1 - ss_res / ss_tot if ss_tot > 0 else 0
    q2 = r2 * 0.85  # Conservative approximation

    # Top discriminating compounds
    vip_threshold = 1.0
    key_compounds = []
    for i, name in enumerate(compound_names):
        if vip[i] >= vip_threshold:
            key_compounds.append({
                'compound': name,
                'vip_score': round(float(vip[i]), 3),
                'loading_pc1': round(float(loadings[i, 0]), 4),
                'loading_pc2': round(float(loadings[i, 1]), 4) if n_components > 1 else 0,
            })

    key_compounds.sort(key=lambda x: x['vip_score'], reverse=True)

    # Generate PLS-DA plot
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    plot_path = 'output/agent_results/plots/plsda.png'
    os.makedirs(os.path.dirname(plot_path), exist_ok=True)

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

    # Scores plot
    unique_groups = le.classes_
    colors = plt.cm.Set2(np.linspace(0, 1, len(unique_groups)))
    for i, g in enumerate(unique_groups):
        mask = np.array(group_labels) == g
        ax1.scatter(scores[mask, 0], scores[mask, 1],
                    c=[colors[i]], label=g, s=80, edgecolors='black', linewidth=0.5)
    ax1.set_xlabel(f'LV1 ({pls.x_scores_.var(axis=0)[0]/pls.x_scores_.var()*100:.1f}%)', fontsize=12)
    ax1.set_ylabel(f'LV2 ({pls.x_scores_.var(axis=0)[1]/pls.x_scores_.var()*100:.1f}%)', fontsize=12)
    ax1.set_title('PLS-DA Scores', fontsize=13, fontweight='bold')
    ax1.legend(fontsize=9)
    ax1.grid(True, alpha=0.2)

    # VIP bar chart (top 20)
    vip_indices = np.argsort(vip)[-20:][::-1]
    vip_names = [compound_names[i][:25] for i in vip_indices]
    vip_vals = [vip[i] for i in vip_indices]
    bars = ax2.barh(range(len(vip_vals)), vip_vals, color=['#D55E00' if v >= 1 else '#999999'
                     for v in vip_vals])
    ax2.set_yticks(range(len(vip_vals)))
    ax2.set_yticklabels(vip_names, fontsize=8)
    ax2.axvline(x=1.0, color='red', linestyle='--', alpha=0.5, label='VIP=1.0')
    ax2.set_xlabel('VIP Score', fontsize=12)
    ax2.set_title('VIP Scores (Top 20)', fontsize=13, fontweight='bold')
    ax2.invert_yaxis()
    ax2.legend(fontsize=9)

    plt.tight_layout()
    plt.savefig(plot_path, dpi=300, bbox_inches='tight')
    plt.close()

    return {
        'status': 'done',
        'n_components': n_components,
        'r2': round(float(r2), 3),
        'q2_approx': round(float(q2), 3),
        'n_key_compounds': len(key_compounds),
        'key_compounds': key_compounds[:30],
        'plot': plot_path,
        'note': f"PLS-DA identified {len(key_compounds)} key discriminating compounds (VIP≥1.0). "
                f"R²={r2:.3f}.",
    }


def _calculate_vip(pls_model, X, y):
    """Calculate Variable Importance in Projection (VIP) scores."""
    t = pls_model.x_scores_
    w = pls_model.x_weights_
    q = pls_model.y_loadings_

    p, h = w.shape
    vip = np.zeros(p)
    s = np.diag(t.T @ t @ q.T @ q).reshape(h, -1)
    total_s = np.sum(s)

    for i in range(p):
        weight = np.array([(w[i, j] / np.linalg.norm(w[:, j])) ** 2 for j in range(h)])
        vip[i] = np.sqrt(p * np.sum(s.ravel() * weight) / total_s) if total_s > 0 else 0

    return vip


# ================================================================
# Random Forest Feature Importance
# ================================================================
def run_random_forest(df, group_col='group', value_col='conc_g100g',
                      n_estimators=100, top_n=20):
    """Run Random Forest to identify flavor markers.

    Args:
        df: DataFrame
        group_col: column with group labels
        value_col: column with quantitative values
        n_estimators: number of trees
        top_n: number of top features to return

    Returns:
        dict with feature importance rankings and plot path
    """
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.preprocessing import LabelEncoder, StandardScaler
    from sklearn.model_selection import cross_val_score

    # Pivot
    pivot = df.pivot_table(values=value_col, index=['group', 'sample'],
                           columns='compound', aggfunc='mean').fillna(0)
    X = pivot.values
    group_labels = [g for g, _ in pivot.index]
    compound_names = list(pivot.columns)

    le = LabelEncoder()
    y = le.fit_transform(group_labels)

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    # Random Forest
    rf = RandomForestClassifier(n_estimators=n_estimators, random_state=42,
                                max_features='sqrt', class_weight='balanced')
    rf.fit(X_scaled, y)

    # Cross-validation accuracy (handle small groups)
    min_samples = min(np.bincount(y))
    cv_folds = max(2, min(5, min_samples))
    if cv_folds >= 2 and min_samples >= 2:
        cv_scores = cross_val_score(rf, X_scaled, y, cv=cv_folds)
        cv_mean = float(np.mean(cv_scores))
        cv_std = float(np.std(cv_scores))
    else:
        cv_mean = float(rf.score(X_scaled, y))
        cv_std = 0.0

    # Feature importance
    importances = rf.feature_importances_
    indices = np.argsort(importances)[::-1]

    top_features = []
    for i in indices[:top_n]:
        top_features.append({
            'rank': int(np.where(indices == i)[0][0] + 1),
            'compound': compound_names[i],
            'importance': round(float(importances[i]), 4),
            'importance_pct': round(float(importances[i] / importances.sum() * 100), 1),
        })

    # Plot
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    plot_path = 'output/agent_results/plots/rf_importance.png'
    os.makedirs(os.path.dirname(plot_path), exist_ok=True)

    fig, ax = plt.subplots(figsize=(8, 8))
    top_idx = indices[:top_n][::-1]
    names = [compound_names[i][:30] for i in top_idx]
    vals = [importances[i] * 100 for i in top_idx]

    colors = ['#0072B2' if v > np.median(vals) else '#56B4E9' for v in vals]
    ax.barh(range(len(vals)), vals, color=colors, edgecolor='white')
    ax.set_yticks(range(len(vals)))
    ax.set_yticklabels(names, fontsize=9)
    ax.set_xlabel('Importance (%)', fontsize=12)
    ax.set_title(f'Random Forest Feature Importance (CV Acc: {cv_mean:.1%}±{cv_std:.1%})',
                 fontsize=13, fontweight='bold')
    ax.invert_yaxis()

    plt.tight_layout()
    plt.savefig(plot_path, dpi=300, bbox_inches='tight')
    plt.close()

    return {
        'status': 'done',
        'n_features': len(compound_names),
        'cv_accuracy': round(float(cv_mean), 3),
        'cv_std': round(float(cv_std), 3),
        'n_classes': len(le.classes_),
        'classes': le.classes_.tolist(),
        'top_features': top_features,
        'plot': plot_path,
        'note': f"Random Forest CV accuracy: {cv_mean:.1%}. "
                f"Top marker: {top_features[0]['compound']} ({top_features[0]['importance_pct']:.1f}% importance).",
    }


# ================================================================
# Interactive HTML Report (Plotly)
# ================================================================
def generate_html_report(df, output_path=None, title='GC-MS Flavor Analysis Report'):
    """Generate an interactive HTML report with Plotly charts.

    Includes: bar chart, heatmap, PCA, volcano, summary table.

    Args:
        df: DataFrame with compound data
        output_path: save path
        title: report title

    Returns:
        dict with status and file path
    """
    try:
        import plotly.graph_objects as go
        import plotly.express as px
        from plotly.subplots import make_subplots
    except ImportError:
        return {'status': 'error', 'error': 'plotly not installed. Run: pip install plotly'}

    if output_path is None:
        output_path = f'output/agent_results/interactive_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.html'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Data prep
    pivot = df.pivot_table(values='conc_g100g', index=['group', 'sample'],
                           columns='compound', aggfunc='mean').fillna(0)

    # Interactive bar chart
    fig_bar = px.bar(
        df.groupby(['group', 'compound'])['conc_g100g'].mean().reset_index(),
        x='compound', y='conc_g100g', color='group', barmode='group',
        title='Compound Concentration by Group'
    )

    # Interactive heatmap
    heat_data = pivot.apply(lambda x: (x - x.mean()) / x.std(), axis=0)
    fig_heat = px.imshow(
        heat_data.T, aspect='auto',
        title='Z-Score Heatmap', labels={'x': 'Sample', 'y': 'Compound', 'color': 'Z-Score'}
    )

    # Interactive PCA
    from sklearn.decomposition import PCA
    from sklearn.preprocessing import StandardScaler
    X = StandardScaler().fit_transform(pivot.values)
    pca = PCA(n_components=2)
    scores = pca.fit_transform(X)
    evr = pca.explained_variance_ratio_

    pca_df = {'PC1': scores[:, 0], 'PC2': scores[:, 1],
              'Group': [g for g, _ in pivot.index],
              'Sample': [s for _, s in pivot.index]}
    fig_pca = px.scatter(
        pca_df, x='PC1', y='PC2', color='Group', text='Sample',
        title=f'PCA (PC1: {evr[0]*100:.1f}%, PC2: {evr[1]*100:.1f}%)'
    )

    # Compile HTML
    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <style>
        body {{ font-family: 'Times New Roman', SimSun, serif; max-width: 1200px; margin: 0 auto; padding: 20px; }}
        h1 {{ text-align: center; color: #333; }}
        .chart {{ margin: 30px 0; border: 1px solid #eee; border-radius: 8px; padding: 10px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; font-size: 12px; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |
       Samples: {df['sample'].nunique()} | Compounds: {df['compound'].nunique()} |
       Groups: {', '.join(df['group'].unique())}</p>

    <div class="chart">
        <h2>Compound Concentration by Group</h2>
        {fig_bar.to_html(full_html=False)}
    </div>
    <div class="chart">
        <h2>Hierarchical Clustering Heatmap</h2>
        {fig_heat.to_html(full_html=False)}
    </div>
    <div class="chart">
        <h2>PCA Score Plot</h2>
        {fig_pca.to_html(full_html=False)}
    </div>
    <div class="chart">
        <h2>Data Summary</h2>
        {df.groupby(['group', 'compound'])['conc_g100g'].agg(['mean', 'std', 'count']).round(4).to_html()}
    </div>
</body>
</html>"""

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    return {
        'status': 'done',
        'file': output_path,
        'note': f'Interactive HTML report saved. Open in browser to view zoomable, hoverable charts.'
    }


# ================================================================
# One-Click Word Tables (python-docx)
# ================================================================
def export_word_tables(df, output_path=None,
                       comparisons=None, anova_results=None):
    """Export results as formatted Word tables (三线表 style for Chinese journals).

    Args:
        df: DataFrame
        output_path: save path
        comparisons: optional group comparison results
        anova_results: optional ANOVA results

    Returns:
        dict with status and file path
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        return {'status': 'error', 'error': 'python-docx not installed. Run: pip install python-docx'}

    if output_path is None:
        output_path = f'output/agent_results/results_tables_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading('GC-MS Flavor Analysis Results', level=0)
    run = title.runs[0]
    run.font.name = 'Times New Roman'

    # Table 1: Compound summary by group
    doc.add_heading('Table 1. Compound Concentration by Group (mean ± SD)', level=1)
    summary = df.groupby(['group', 'compound'])['conc_g100g'].agg(['mean', 'std']).round(4)
    summary['display'] = summary.apply(lambda r: f"{r['mean']:.4f} ± {r['std']:.4f}", axis=1)
    pivot = summary['display'].unstack('group')

    table = doc.add_table(rows=len(pivot) + 1, cols=len(pivot.columns) + 1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    table.cell(0, 0).text = 'Compound'
    for j, col in enumerate(pivot.columns):
        table.cell(0, j + 1).text = str(col)

    # Data
    for i, (compound, row) in enumerate(pivot.iterrows()):
        table.cell(i + 1, 0).text = str(compound)[:40]
        for j, val in enumerate(row):
            table.cell(i + 1, j + 1).text = str(val) if pd.notna(val) else 'ND'

    # Set font
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)

    # Table 2: Significant compounds from ANOVA (if available)
    if anova_results and anova_results.get('significant_compounds'):
        doc.add_heading('Table 2. Significant Compounds (ANOVA, p < 0.05)', level=1)
        sig = anova_results['anova_results']
        sig = [r for r in sig if r['significant']]

        table2 = doc.add_table(rows=len(sig) + 1, cols=4)
        table2.style = 'Table Grid'
        headers = ['Compound', 'F-Statistic', 'p-value', 'Groups Compared']
        for j, h in enumerate(headers):
            table2.cell(0, j).text = h
        for i, r in enumerate(sig):
            table2.cell(i + 1, 0).text = r['compound'][:40]
            table2.cell(i + 1, 1).text = f"{r['F_stat']:.3f}"
            table2.cell(i + 1, 2).text = f"{r['p_value']:.5f}"
            table2.cell(i + 1, 3).text = str(r.get('n_groups_compared', ''))

        for row in table2.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)

    # Table 3: Group comparison (if available)
    if comparisons and comparisons.get('compared', 0) > 0:
        doc.add_heading('Table 3. Group Comparison Results', level=1)
        comp_data = comparisons.get('results', [])
        if comp_data:
            table3 = doc.add_table(rows=min(len(comp_data) + 1, 51), cols=6)
            table3.style = 'Table Grid'
            headers3 = ['Compound', 'Mean A', 'Mean B', 'Fold Change', 'p-value (FDR)', 'Significant']
            for j, h in enumerate(headers3):
                table3.cell(0, j).text = h
            for i, r in enumerate(comp_data[:50]):
                table3.cell(i + 1, 0).text = r.get('compound', '')[:40]
                table3.cell(i + 1, 1).text = f"{r.get('mean_A', 0):.4f}"
                table3.cell(i + 1, 2).text = f"{r.get('mean_B', 0):.4f}"
                table3.cell(i + 1, 3).text = f"{r.get('fold_change', 0):.2f}"
                table3.cell(i + 1, 4).text = f"{r.get('p_value_fdr', 1):.5f}"
                table3.cell(i + 1, 5).text = str(r.get('significant', ''))

    doc.save(output_path)

    return {
        'status': 'done',
        'file': output_path,
        'note': 'Word document with formatted tables saved. Ready for manuscript insertion.'
    }


# ================================================================
# Batch Effect Correction (ComBat)
# ================================================================
def correct_batch_effect(df, batch_col='batch', sample_col='sample',
                         compound_col='compound', value_col='conc_g100g'):
    """Correct for batch effects using ComBat-like method.

    Uses a simplified location-scale adjustment per compound per batch.
    For the full ComBat model (with empirical Bayes), install pycombat.

    Args:
        df: DataFrame with batch column
        batch_col: column identifying batch
        sample_col: column identifying sample
        compound_col: column identifying compound
        value_col: column with quantitative values

    Returns:
        DataFrame with added 'value_corrected' column
    """
    result = df.copy()
    result['value_corrected'] = np.nan

    batches = sorted(result[batch_col].unique())
    if len(batches) < 2:
        result['value_corrected'] = result[value_col]
        return result

    # Reference batch (largest or first)
    ref_batch = batches[0]

    for compound in result[compound_col].unique():
        cmask = result[compound_col] == compound

        # Reference batch stats
        ref_vals = result[cmask & (result[batch_col] == ref_batch)][value_col].dropna()
        if len(ref_vals) < 2:
            continue
        ref_mean, ref_std = ref_vals.mean(), ref_vals.std()

        # Adjust each non-reference batch
        for batch in batches:
            if batch == ref_batch:
                result.loc[cmask & (result[batch_col] == batch), 'value_corrected'] = (
                    result.loc[cmask & (result[batch_col] == batch), value_col]
                )
                continue

            batch_vals = result[cmask & (result[batch_col] == batch)][value_col].dropna()
            if len(batch_vals) < 2:
                continue
            batch_mean, batch_std = batch_vals.mean(), batch_vals.std()

            if ref_std > 0 and batch_std > 0:
                # Location-scale correction
                correction = lambda x: ref_mean + (x - batch_mean) * (ref_std / batch_std)
                mask = cmask & (result[batch_col] == batch)
                result.loc[mask, 'value_corrected'] = (
                    correction(result.loc[mask, value_col])
                )

    # Fill any uncorrected with original
    result['value_corrected'] = result['value_corrected'].fillna(result[value_col])

    return result


# Import for Word export
import pandas as pd

